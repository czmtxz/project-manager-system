# -*- coding: utf-8 -*-
"""
OCR识别工具模块
基于 RapidOCR-ONNXRuntime 实现本地OCR识别
支持：增值税发票、过磅单、运单、通用文字识别
"""

import os
import re
import json
from datetime import datetime

# OCR引擎初始化（延迟加载）
_ocr_engine = None


def get_ocr_engine():
    """获取OCR引擎（单例模式）"""
    global _ocr_engine
    if _ocr_engine is None:
        try:
            from rapidocr_onnxruntime import RapidOCR
            _ocr_engine = RapidOCR()
        except ImportError:
            raise ImportError("请安装 rapidocr-onnxruntime: pip install rapidocr-onnxruntime")
    return _ocr_engine


def ocr_image(image_path):
    """
    对图片进行OCR识别，返回所有文字及其位置
    :param image_path: 图片路径
    :return: list of (text, confidence, box)
    """
    engine = get_ocr_engine()
    result, _ = engine(image_path)
    
    if not result:
        return []
    
    # result 格式: [[box, text, confidence], ...]
    return [(item[1], item[2], item[0]) for item in result]


def ocr_image_to_lines(image_path):
    """
    对图片进行OCR识别，返回按行排列的文字
    :param image_path: 图片路径
    :return: list of str
    """
    results = ocr_image(image_path)
    # 按Y坐标排序（从上到下）
    sorted_results = sorted(results, key=lambda x: x[2][0][1])
    return [item[0] for item in sorted_results]


# ==================== 增值税发票识别 ====================

def recognize_invoice(image_path):
    """
    识别增值税发票图片
    :param image_path: 发票图片路径
    :return: dict 发票信息
    """
    lines = ocr_image_to_lines(image_path)
    full_text = '\n'.join(lines)
    
    result = {
        'invoice_no': '',       # 发票号码
        'invoice_code': '',     # 发票代码
        'invoice_date': '',     # 开票日期
        'amount': 0,            # 金额
        'tax_rate': 0,          # 税率
        'tax_amount': 0,        # 税额
        'total_amount': 0,      # 价税合计
        'seller': '',           # 销方名称
        'buyer': '',            # 购方名称
        'invoice_type': '',     # 发票类型
        'raw_text': full_text,
        'confidence': 0
    }
    
    # 识别发票类型
    if '增值税' in full_text and '专用发票' in full_text:
        result['invoice_type'] = '增值税专用发票'
    elif '增值税' in full_text and '普通发票' in full_text:
        result['invoice_type'] = '增值税普通发票'
    elif '机动车' in full_text:
        result['invoice_type'] = '机动车销售统一发票'
    elif '电子发票' in full_text:
        result['invoice_type'] = '电子发票'
    
    # 识别发票号码（8位或10位数字）
    invoice_no_match = re.search(r'发票号码[：:]\s*(\d{8,20})', full_text)
    if not invoice_no_match:
        invoice_no_match = re.search(r'No[.:：]\s*(\d{8,20})', full_text)
    if invoice_no_match:
        result['invoice_no'] = invoice_no_match.group(1)
    
    # 识别发票代码
    code_match = re.search(r'发票代码[：:]\s*(\d{10,12})', full_text)
    if code_match:
        result['invoice_code'] = code_match.group(1)
    
    # 识别开票日期
    date_match = re.search(r'开票日期[：:]\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', full_text)
    if date_match:
        result['invoice_date'] = f"{date_match.group(1)}-{int(date_match.group(2)):02d}-{int(date_match.group(3)):02d}"
    else:
        date_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', full_text)
        if date_match:
            result['invoice_date'] = f"{date_match.group(1)}-{int(date_match.group(2)):02d}-{int(date_match.group(3)):02d}"
    
    # 识别金额（多种模式）
    # 价税合计
    total_match = re.search(r'[价合][税计][合总][：:\s]*[¥￥]?\s*([0-9,]+\.?\d*)', full_text)
    if total_match:
        result['total_amount'] = float(total_match.group(1).replace(',', ''))
    
    # 金额
    amount_match = re.search(r'(?:金额|不含税金额)[：:\s]*[¥￥]?\s*([0-9,]+\.?\d*)', full_text)
    if amount_match:
        result['amount'] = float(amount_match.group(1).replace(',', ''))
    
    # 税额
    tax_match = re.search(r'税[额率][：:\s]*[¥￥]?\s*([0-9,]+\.?\d*)', full_text)
    if tax_match:
        val = float(tax_match.group(1).replace(',', ''))
        if val < 1:  # 可能是税率
            result['tax_rate'] = val * 100
        else:
            result['tax_amount'] = val
    
    # 税率
    rate_match = re.search(r'税率[：:\s]*(\d+)%', full_text)
    if rate_match:
        result['tax_rate'] = float(rate_match.group(1))
    
    # 如果有金额和税额但没有价税合计
    if result['amount'] > 0 and result['tax_amount'] > 0 and result['total_amount'] == 0:
        result['total_amount'] = round(result['amount'] + result['tax_amount'], 2)
    
    # 如果有金额和税率但没有税额
    if result['amount'] > 0 and result['tax_rate'] > 0 and result['tax_amount'] == 0:
        result['tax_amount'] = round(result['amount'] * result['tax_rate'] / 100, 2)
        result['total_amount'] = round(result['amount'] + result['tax_amount'], 2)
    
    # 识别销方名称
    seller_match = re.search(r'销[方售][名称]*[：:\s]*([^\n]+)', full_text)
    if seller_match:
        result['seller'] = seller_match.group(1).strip()
    
    # 识别购方名称
    buyer_match = re.search(r'购[方买][名称]*[：:\s]*([^\n]+)', full_text)
    if buyer_match:
        result['buyer'] = buyer_match.group(1).strip()
    
    # 计算置信度
    fields_found = sum(1 for k in ['invoice_no', 'amount', 'invoice_date'] if result[k])
    result['confidence'] = fields_found / 3
    
    return result


# ==================== 过磅单/运单识别 ====================

def recognize_weighbridge(image_path):
    """
    识别过磅单图片
    :param image_path: 过磅单图片路径
    :return: dict 过磅信息
    """
    lines = ocr_image_to_lines(image_path)
    full_text = '\n'.join(lines)
    
    result = {
        'vehicle_no': '',       # 车牌号
        'driver_name': '',      # 司机姓名
        'transport_date': '',   # 运输日期
        'material_name': '',    # 材料名称
        'gross_weight': 0,      # 毛重
        'tare_weight': 0,       # 皮重
        'net_weight': 0,        # 净重
        'unit': '吨',           # 单位
        'raw_text': full_text,
        'type': 'weighbridge',
        'confidence': 0
    }
    
    # 识别车牌号（多种格式）
    plate_patterns = [
        r'[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤川青藏琼宁][A-Z][A-Z0-9]{5,6}',
        r'[A-Z][A-Z0-9]{5,7}',
    ]
    for pattern in plate_patterns:
        plate_match = re.search(pattern, full_text)
        if plate_match:
            result['vehicle_no'] = plate_match.group(0)
            break
    
    # 识别日期
    date_match = re.search(r'(\d{4})[年/\-.](\d{1,2})[月/\-.](\d{1,2})[日]?', full_text)
    if date_match:
        result['transport_date'] = f"{date_match.group(1)}-{int(date_match.group(2)):02d}-{int(date_match.group(3)):02d}"
    
    # 识别重量
    # 毛重
    gross_match = re.search(r'毛\s*重[：:\s]*([0-9,]+\.?\d*)', full_text)
    if not gross_match:
        gross_match = re.search(r'进场重量[：:\s]*([0-9,]+\.?\d*)', full_text)
    if gross_match:
        result['gross_weight'] = float(gross_match.group(1).replace(',', ''))
    
    # 皮重
    tare_match = re.search(r'皮\s*重[：:\s]*([0-9,]+\.?\d*)', full_text)
    if not tare_match:
        tare_match = re.search(r'空车重量[：:\s]*([0-9,]+\.?\d*)', full_text)
    if tare_match:
        result['tare_weight'] = float(tare_match.group(1).replace(',', ''))
    
    # 净重
    net_match = re.search(r'净\s*重[：:\s]*([0-9,]+\.?\d*)', full_text)
    if not net_match:
        net_match = re.search(r'实际重量[：:\s]*([0-9,]+\.?\d*)', full_text)
    if net_match:
        result['net_weight'] = float(net_match.group(1).replace(',', ''))
    
    # 如果有毛重和皮重但没有净重
    if result['gross_weight'] > 0 and result['tare_weight'] > 0 and result['net_weight'] == 0:
        result['net_weight'] = round(result['gross_weight'] - result['tare_weight'], 2)
    
    # 识别材料名称
    material_patterns = [
        r'(砂石|碎石|砂子|石子|水泥|钢材|螺纹钢|河砂|机制砂|石粉|瓜子片|商品混凝土|沥青)',
        r'材料[名称]*[：:\s]*([^\s\n]+)',
        r'品名[：:\s]*([^\s\n]+)',
        r'货物[名称]*[：:\s]*([^\s\n]+)',
    ]
    for pattern in material_patterns:
        mat_match = re.search(pattern, full_text)
        if mat_match:
            result['material_name'] = mat_match.group(1).strip()
            break
    
    # 识别司机
    driver_match = re.search(r'(?:司机|驾驶员)[：:\s]*([^\s\n]{2,4})', full_text)
    if driver_match:
        result['driver_name'] = driver_match.group(1).strip()
    
    # 计算置信度
    fields_found = sum(1 for k in ['net_weight', 'vehicle_no', 'transport_date'] if result[k])
    result['confidence'] = fields_found / 3
    
    return result


def recognize_freight_invoice(image_path):
    """
    识别运费发票图片
    :param image_path: 运费发票图片路径
    :return: dict 运费发票信息
    """
    lines = ocr_image_to_lines(image_path)
    full_text = '\n'.join(lines)
    
    result = {
        'invoice_no': '',
        'freight_amount': 0,
        'vehicle_no': '',
        'transport_date': '',
        'route': '',           # 运输路线
        'raw_text': full_text,
        'type': 'freight_invoice',
        'confidence': 0
    }
    
    # 先尝试作为普通发票识别
    invoice_data = recognize_invoice(image_path)
    result['invoice_no'] = invoice_data.get('invoice_no', '')
    result['freight_amount'] = invoice_data.get('total_amount', 0) or invoice_data.get('amount', 0)
    result['transport_date'] = invoice_data.get('invoice_date', '')
    
    # 识别车牌号
    plate_match = re.search(r'[京津沪渝冀豫云辽黑湘皖鲁新苏浙赣鄂桂甘晋蒙陕吉闽贵粤川青藏琼宁][A-Z][A-Z0-9]{5,6}', full_text)
    if plate_match:
        result['vehicle_no'] = plate_match.group(0)
    
    # 识别运输路线
    route_match = re.search(r'([^\s]+)\s*[→到至\-]\s*([^\s]+)', full_text)
    if route_match:
        result['route'] = f"{route_match.group(1)}→{route_match.group(2)}"
    
    # 如果金额为0，尝试其他模式
    if result['freight_amount'] == 0:
        amount_match = re.search(r'[运费][：:\s]*[¥￥]?\s*([0-9,]+\.?\d*)', full_text)
        if amount_match:
            result['freight_amount'] = float(amount_match.group(1).replace(',', ''))
    
    fields_found = sum(1 for k in ['freight_amount', 'vehicle_no'] if result[k])
    result['confidence'] = fields_found / 2
    
    return result


def recognize_transport_image(image_path):
    """
    自动识别运输相关图片（自动判断类型）
    :param image_path: 图片路径
    :return: dict 识别结果
    """
    lines = ocr_image_to_lines(image_path)
    full_text = '\n'.join(lines)
    
    # 判断图片类型
    if '增值税' in full_text or '发票' in full_text:
        if '运' in full_text or '运输' in full_text or ' freight' in full_text.lower():
            return recognize_freight_invoice(image_path)
        else:
            return recognize_invoice(image_path)
    elif '毛重' in full_text or '皮重' in full_text or '净重' in full_text or '过磅' in full_text:
        return recognize_weighbridge(image_path)
    elif '磅单' in full_text or '运单' in full_text:
        return recognize_weighbridge(image_path)
    else:
        # 默认尝试过磅单识别
        return recognize_weighbridge(image_path)


# ==================== Word文档解析 ====================

def parse_contract_docx(docx_path):
    """
    解析合同Word文档，提取关键字段
    :param docx_path: Word文档路径
    :return: dict 合同信息
    """
    from docx import Document
    
    doc = Document(docx_path)
    
    # 提取所有段落文本
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text = '\n'.join(paragraphs)
    
    # 提取表格文本
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_texts.append(' | '.join(row_text))
    table_text = '\n'.join(table_texts)
    combined_text = full_text + '\n' + table_text
    
    result = {
        'contract_no': '',
        'contract_name': '',
        'contract_type': '',
        'party': '',            # 对方单位
        'amount': 0,
        'tax_rate': 0,
        'sign_date': '',
        'start_date': '',
        'end_date': '',
        'description': '',
        'raw_text': combined_text,
        'confidence': 0
    }
    
    # 识别合同编号
    no_patterns = [
        r'合同编号[：:\s]*([A-Z0-9\-]+)',
        r'编号[：:\s]*([A-Z0-9\-]+)',
        r'(HT-\d{4}-\d+)',
        r'(合同编号[：:\s]*[^\s\n]+)',
    ]
    for pattern in no_patterns:
        match = re.search(pattern, combined_text)
        if match:
            result['contract_no'] = match.group(1).strip()
            break
    
    # 识别合同名称（通常在标题位置）
    for p in paragraphs[:5]:
        if '合同' in p and len(p) < 50:
            result['contract_name'] = p.strip()
            break
    
    # 识别合同类型
    if '采购' in combined_text:
        result['contract_type'] = '支出合同'
    elif '销售' in combined_text:
        result['contract_type'] = '收入合同'
    elif '工程' in combined_text:
        result['contract_type'] = '支出合同'
    elif '运输' in combined_text:
        result['contract_type'] = '支出合同'
    elif '服务' in combined_text:
        result['contract_type'] = '支出合同'
    
    # 识别对方单位（甲方/乙方）
    party_patterns = [
        r'(?:甲方|供方|卖方)[：:\s]*([^\n,，]+)',
        r'(?:乙方|需方|买方)[：:\s]*([^\n,，]+)',
    ]
    for pattern in party_patterns:
        match = re.search(pattern, combined_text)
        if match:
            result['party'] = match.group(1).strip()
            break
    
    # 识别金额
    amount_patterns = [
        r'合同总[价金额][：:\s]*[¥￥]?\s*([0-9,]+\.?\d*)',
        r'总[价金额][：:\s]*[¥￥]?\s*([0-9,]+\.?\d*)\s*[万]?元',
        r'[¥￥]\s*([0-9,]+\.?\d*)',
        r'人民币[：:\s]*([0-9,]+\.?\d*)\s*元',
    ]
    for pattern in amount_patterns:
        match = re.search(pattern, combined_text)
        if match:
            val = float(match.group(1).replace(',', ''))
            # 如果文本中有"万"字，需要乘以10000
            if '万' in match.group(0):
                val *= 10000
            result['amount'] = round(val, 2)
            break
    
    # 识别税率
    rate_match = re.search(r'税率[：:\s]*(\d+)%', combined_text)
    if rate_match:
        result['tax_rate'] = float(rate_match.group(1))
    
    # 识别签订日期
    date_patterns = [
        r'签订日期[：:\s]*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
        r'签署日期[：:\s]*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
        r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
    ]
    for pattern in date_patterns:
        match = re.search(pattern, combined_text)
        if match:
            result['sign_date'] = f"{match.group(1)}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"
            break
    
    # 识别合同期限
    period_match = re.search(r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*[至到]\s*(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', combined_text)
    if period_match:
        result['start_date'] = f"{period_match.group(1)}-{int(period_match.group(2)):02d}-{int(period_match.group(3)):02d}"
        result['end_date'] = f"{period_match.group(4)}-{int(period_match.group(5)):02d}-{int(period_match.group(6)):02d}"
    
    # 生成描述（取前几段非空文本）
    desc_parts = [p for p in paragraphs if len(p) > 10 and '合同' not in p and '编号' not in p][:3]
    result['description'] = '；'.join(desc_parts) if desc_parts else ''
    
    # 计算置信度
    fields_found = sum(1 for k in ['contract_no', 'amount', 'party'] if result[k])
    result['confidence'] = fields_found / 3
    
    return result
